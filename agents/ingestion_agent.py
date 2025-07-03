# ingestion_agent.py
import os
import logging
from typing import List, Dict, Any

# No need for sys.path manipulation here if main.py handles it
from mcp.mcp_message_protocol import MCPMessage, MCPMessageType, BaseMCPMessageHandler
from utils.utils_document_processor import DocumentProcessor
from utils.utils_vector_store import VectorStore

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class IngestionAgent(BaseMCPMessageHandler):
    """
    The IngestionAgent is responsible for processing uploaded documents:
    - Parsing documents and extracting raw text.
    - Splitting extracted text into manageable chunks.
    - Storing these chunks and their metadata (including source) in the vector store.
    """
    def __init__(self, vector_store: VectorStore):
        super().__init__("IngestionAgent")
        self.document_processor = DocumentProcessor()
        self.vector_store = vector_store
        logging.info("IngestionAgent initialized.")

    def process_message(self, message: MCPMessage) -> MCPMessage:
        """
        Processes incoming MCP messages, specifically handling DOCUMENT_UPLOAD messages.

        Args:
            message (MCPMessage): The incoming message.

        Returns:
            MCPMessage: A response message indicating success or failure.
        """
        if message.type == MCPMessageType.DOCUMENT_UPLOAD:
            return self._handle_document_upload(message)
        else:
            logging.warning(f"IngestionAgent received unhandled message type: {message.type.name}")
            return self._create_response_message(
                receiver=message.sender,
                message_type=MCPMessageType.ERROR,
                payload={"error": f"Unhandled message type: {message.type.name}"},
                trace_id=message.trace_id
            )

    def _handle_document_upload(self, message: MCPMessage) -> MCPMessage:
        """
        Handles the DOCUMENT_UPLOAD message by processing each file.
        """
        files_info: List[Dict[str, Any]] = message.payload.get("files", [])
        if not files_info:
            logging.warning("No files found in DOCUMENT_UPLOAD payload.")
            return self._create_response_message(
                receiver=message.sender,
                message_type=MCPMessageType.ERROR,
                payload={"error": "No files provided for upload."},
                trace_id=message.trace_id
            )

        processed_files_count = 0
        failed_files_count = 0
        processed_documents_metadata = []

        for file_info in files_info:
            file_name = file_info.get("name")
            file_path = file_info.get("path") # In a real app, this would be a temporary file path
            original_file_name = file_info.get("original_name", file_name) # For UI display

            if not file_name or not file_path or not os.path.exists(file_path):
                logging.error(f"Invalid file info or file not found for: {file_info}")
                failed_files_count += 1
                continue

            logging.info(f"Processing document: {original_file_name} from {file_path}")
            try:
                extracted_text, doc_type = self.document_processor.parse_document(file_path)

                if not extracted_text:
                    logging.warning(f"No text extracted from {original_file_name} ({doc_type}). Skipping.")
                    failed_files_count += 1
                    continue

                chunks = self.document_processor.split_text_into_chunks(
                    extracted_text, chunk_size=1000, chunk_overlap=100
                )
                logging.info(f"Split {original_file_name} into {len(chunks)} chunks.")

                # Prepare metadata for each chunk
                chunk_metadatas = []
                for i, chunk in enumerate(chunks):
                    chunk_metadatas.append({
                        "text": chunk,
                        "source": original_file_name,
                        "chunk_index": i,
                        "doc_type": doc_type,
                        # Add other relevant metadata like page number if available from parser
                    })

                self.vector_store.add_texts(chunks, chunk_metadatas)
                processed_files_count += 1
                processed_documents_metadata.append({
                    "name": original_file_name,
                    "chunks_count": len(chunks),
                    "status": "success"
                })
            except Exception as e:
                logging.error(f"Error processing document {original_file_name}: {e}", exc_info=True)
                failed_files_count += 1
                processed_documents_metadata.append({
                    "name": original_file_name,
                    "status": "failed",
                    "error": str(e)
                })

        response_payload = {
            "status": "success" if failed_files_count == 0 else "partial_success" if processed_files_count > 0 else "failure",
            "processed_files_count": processed_files_count,
            "failed_files_count": failed_files_count,
            "processed_documents_metadata": processed_documents_metadata,
            "message": f"Processed {processed_files_count} files, failed {failed_files_count}."
        }

        return self._create_response_message(
            receiver=message.sender,
            message_type=MCPMessageType.LLM_RESPONSE, # Or a dedicated INGESTION_COMPLETE type
            payload=response_payload,
            trace_id=message.trace_id
        )

# Example Usage (for testing)
if __name__ == "__main__":
    # For standalone testing, you might need to adjust sys.path here if not run from project root
    import sys
    import os
    current_dir_ia_test = os.path.dirname(os.path.abspath(__file__))
    project_root_ia_test = os.path.abspath(os.path.join(current_dir_ia_test, '..')) # One level up to 'AgenticRAGChatbot' assuming 'agents' is directly under it
    if project_root_ia_test not in sys.path:
        sys.path.insert(0, project_root_ia_test)
        logging.info(f"Added {project_root_ia_test} to sys.path for standalone testing.")

    # In a real scenario, you'd integrate this with Streamlit file uploads
    # For testing, we'll create dummy files and a dummy vector store
    import tempfile
    import os

    # Create dummy files
    dummy_files = []
    temp_dir = tempfile.mkdtemp()

    with open(os.path.join(temp_dir, "test.txt"), "w") as f:
        f.write("This is a test text file. It contains some basic information for chunking.")
    dummy_files.append({"name": "test.txt", "path": os.path.join(temp_dir, "test.txt"), "original_name": "test.txt"})

    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("This is a Word document. It has multiple paragraphs.")
        doc.add_paragraph("Second paragraph for testing docx parsing.")
        docx_path = os.path.join(temp_dir, "report.docx")
        doc.save(docx_path)
        dummy_files.append({"name": "report.docx", "path": docx_path, "original_name": "report.docx"})
    except ImportError:
        print("python-docx not installed, skipping DOCX dummy file creation.")
        logging.warning("python-docx not installed, skipping DOCX dummy file creation for test.")


    # Setup VectorStore (using the actual VectorStore for this example)
    vs = VectorStore()
    ingestion_agent = IngestionAgent(vs)

    # Simulate a DOCUMENT_UPLOAD message
    upload_message = MCPMessage(
        sender="UI",
        receiver="IngestionAgent",
        message_type=MCPMessageType.DOCUMENT_UPLOAD,
        payload={"files": dummy_files}
    )

    print(f"Simulating document upload message: {upload_message.to_json()}")
    response_message = ingestion_agent.process_message(upload_message)
    print(f"\nIngestionAgent Response: {response_message.to_json()}")

    print(f"\nTotal items in vector store after ingestion: {vs.index.ntotal}")

    # Clean up dummy files
    for f_info in dummy_files:
        try:
            if os.path.exists(f_info["path"]):
                os.remove(f_info["path"])
        except OSError as e:
            print(f"Error cleaning up {f_info['path']}: {e}")
            logging.error(f"Error cleaning up {f_info['path']}: {e}")
    os.rmdir(temp_dir)